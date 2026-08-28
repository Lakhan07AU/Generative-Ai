"""Security Policy RAG (Part 2).

Pipeline:  Document (PDF/DOCX/TXT) -> text extraction -> section-aware chunking
with overlap -> embeddings -> Qdrant `policy_chunks`.

Also implements policy questions: retrieve the relevant policy, retrieve video
evidence for the same question, compare observed behaviour with the policy, and
produce a POLICY-ASSESSED finding. The policy is never invented - if no policy
chunk is retrieved the system returns UNKNOWN.
"""

from __future__ import annotations

import logging
import os
import re
import tempfile
from typing import Any, Optional

from app.core.config import settings
from app.database import models
from app.database.session import SessionLocal
from app.ai import provider
from app.ai.embeddings import embeddings
from app.ai.qdrant_service import qdrant
from app.rag.video_rag import analyze_query, retrieve_evidence, evidence_cards

logger = logging.getLogger(__name__)

CHUNK_SIZE = 600
CHUNK_OVERLAP = 120

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(filename: str, data: bytes) -> tuple[str, str]:
    """Return (source_format, text) for a policy document."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return "pdf", _extract_pdf(data)
    if ext == ".docx":
        return "docx", _extract_docx(data)
    if ext == ".txt":
        return "txt", data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type '{ext}'. Allowed: {sorted(ALLOWED_EXTENSIONS)}")


def _extract_pdf(data: bytes) -> str:
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text)


# ---------------------------------------------------------------------------
# Section-aware chunking
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(section|article|policy|rule|part|chapter)\s+[0-9IVX]+\.?", re.IGNORECASE)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[dict]:
    """Split text into chunks, trying to respect section boundaries with overlap."""
    if not text.strip():
        return []
    blocks = [b.strip() for b in text.splitlines() if b.strip()]
    # list-like lines (bullets/numbered) are kept together-ish; simple join
    current = ""
    current_section = "General"
    chunks: list[dict] = []
    chunk_index = 0

    def _flush():
        nonlocal current, chunk_index
        if current.strip():
            chunks.append(
                {
                    "section": current_section,
                    "page": 1,
                    "chunk_index": chunk_index,
                    "text": current.strip(),
                }
            )
            chunk_index += 1
        current = ""

    for block in blocks:
        if _HEADING_RE.match(block) and len(block) < 120:
            _flush()
            current_section = block
            current = block + "\n"
            continue
        if len(current) + len(block) <= chunk_size:
            current += (("\n" if current else "") + block)
        else:
            _flush()
            # overlap: prepend tail of previous chunk
            tail = " ".join(current.split()[-10:])
            current = (tail + " " + block) if tail else block
    _flush()
    return chunks


# ---------------------------------------------------------------------------
# Ingest
# ---------------------------------------------------------------------------

def ingest_policy(filename: str, data: bytes, user_id: int | None = None, storage_path: Optional[str] = None) -> dict:
    """Extract, chunk, embed and index a policy document into Qdrant + DB."""
    source_format, text = extract_text(filename, data)
    if not text.strip():
        raise ValueError("No extractable text found in the document")

    db = SessionLocal()
    try:
        existing_policy_id = _next_policy_id(db)
        policy = models.PolicyDocument(
            policy_id=existing_policy_id,
            document_name=os.path.splitext(filename)[0],
            filename=filename,
            storage_path=storage_path,
            source_format=source_format,
            uploaded_by_user_id=user_id,
            status="INDEXED",
        )
        db.add(policy)
        db.flush()
        db.commit()
        db.refresh(policy)

        chunks = chunk_text(text)
        qdrant.ensure_collections()
        vectors = embeddings.embed_texts([c["text"] for c in chunks])
        for chunk, vec in zip(chunks, vectors):
            pc = models.PolicyChunk(
                policy_id=policy.id,
                section=chunk["section"],
                page=chunk["page"],
                chunk_index=chunk["chunk_index"],
                text=chunk["text"],
            )
            db.add(pc)
            db.flush()
            qdrant.index_policy(pc.id, vec, {
                "policy_id": policy.id,
                "document_name": policy.document_name,
                "section": chunk["section"],
                "page": chunk["page"],
                "chunk_index": chunk["chunk_index"],
                "text": chunk["text"],
            })
        db.commit()
        from app.audit.service import record_audit

        record_audit(
            db,
            "policy_upload",
            user_id=user_id,
            entity_type="policy",
            entity_id=policy.id,
            details=f"policy_id={policy.policy_id} chunks={len(chunks)}",
        )
        return {"policy_id": policy.policy_id, "document_name": policy.document_name, "chunks": len(chunks)}
    finally:
        db.close()


def _next_policy_id(db) -> str:
    last = db.query(models.PolicyDocument).order_by(models.PolicyDocument.id.desc()).first()
    if last and last.policy_id.startswith("POL-"):
        try:
            return f"POL-{int(last.policy_id.split('-')[1]) + 1:04d}"
        except (ValueError, IndexError):
            pass
    return "POL-0001"


# ---------------------------------------------------------------------------
# Retrieval
# ---------------------------------------------------------------------------

def search_policies(query: str, limit: int = 5) -> list[dict]:
    vec = embeddings.embed_text(query)
    hits = qdrant.search_policy(vec, limit)
    return [
        {
            "score": round(float(h["score"]), 4),
            "policy_id": h["payload"].get("policy_id"),
            "document_name": h["payload"].get("document_name"),
            "section": h["payload"].get("section"),
            "page": h["payload"].get("page"),
            "chunk_index": h["payload"].get("chunk_index"),
            "text": h["payload"].get("text"),
        }
        for h in hits
    ]


# ---------------------------------------------------------------------------
# Policy question -> POLICY-ASSESSED finding
# ---------------------------------------------------------------------------

def policy_question(db, question: str, video_id: Optional[int] = None) -> dict:
    """Assess a question against the security policy and available video evidence."""
    # Retrieve policy (never invented: if none retrieved, we return UNKNOWN)
    policies = search_policies(question, limit=settings.RAG_RERANK_KEEP)
    # Retrieve video evidence
    filters = {"video_id": video_id} if video_id else None
    hits = retrieve_evidence(db, question, filters)
    analysis = analyze_query(question)
    cards = evidence_cards(db, hits, analysis)
    verified_cards = [c for c in cards if c["verification"]["verified"]]

    finding_status = "UNKNOWN"
    description_parts = []
    policy_used = None

    if not policies:
        finding_status = "UNKNOWN"
        description = "UNKNOWN - no policy was retrieved for this question. A policy assessment cannot be made without a policy source."
        return _policy_payload(db, question, finding_status, description, [], verified_cards, None)

    policy_used = policies[0]
    description_parts.append(f"Policy section '{policies[0]['section']}' applies: {policies[0]['text'][:400]}")

    if not verified_cards:
        # policy found but no observed behaviour -> cannot assess
        description = (
            "UNKNOWN - INSUFFICIENT EVIDENCE. A policy section was found, but no observed "
            "video evidence matched the question, so consistency cannot be assessed. "
            f"{policy_used['document_name']} sec. {policy_used['section']}"
        )
        if provider.available():
            desc = provider.generate_text(
                f"Summarise consistency between policy: {policies[0]['text'][:400]} and evidence-set: none found. "
                "State it cannot be assessed."
            )
            description = desc + f"\n[Policy: {policies[0]['document_name']} sec. {policies[0]['section']}]"
        return _policy_payload(db, question, finding_status, description, policies, verified_cards, policy_used)

    # policy + observed evidence -> POLICY-ASSESSED
    finding_status = "POLICY-ASSESSED"
    obs = verified_cards[0]
    description_parts.append(
        f"Observed behaviour: {obs['clip_public_id']} at {obs['timestamp']:.1f}s - {obs['description']}"
    )
    if provider.available():
        combined = provider.generate_text(
            f"Compare observed behaviour with the security policy and state whether it is (CONSISTENT / NOT CONSISTENT / CANNOT DETERMINE).\n"
            f"Policy: {policies[0]['text'][:500]}\nObserved: {obs['description']}"
        )
        description_parts.append(combined)
    else:
        description_parts.append(
            "[SIMULATED] Consistency determination is deferred to a configured reviewer model. "
            "Observed behaviour and the retrieved policy section are both recorded."
        )
    description = "\n".join(description_parts)
    return _policy_payload(db, question, finding_status, description, policies, verified_cards, policy_used)


def _policy_payload(db, question, finding_status, description, policies, cards, policy_used) -> dict:
    video_id = cards[0]["video_id"] if cards else None
    clip_id = cards[0]["clip_id"] if cards else None
    return {
        "question": question,
        "status": finding_status,
        "description": description,
        "policy_sections": policies,
        "evidence": cards,
        "video_id": video_id,
        "clip_id": clip_id,
        "policy_id": policy_used["policy_id"] if policy_used else None,
    }
